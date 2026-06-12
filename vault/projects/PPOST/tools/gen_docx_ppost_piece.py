# -*- coding: utf-8 -*-
"""
PPOST 件數比對邏輯修正 - DOCX 產生器 (make-docx Pixiu style)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

# 頁面設定
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# 顏色
C_TITLE  = RGBColor(0x1F, 0x49, 0x7D)
C_H1     = RGBColor(0x2E, 0x75, 0xB6)
C_H2     = RGBColor(0x2E, 0x75, 0xB6)
C_H3     = RGBColor(0x40, 0x40, 0x40)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_ORANGE = RGBColor(0xE2, 0x6B, 0x0A)
C_GREEN  = RGBColor(0x37, 0x86, 0x10)

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*rgb))
    tcPr.append(shd)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = C_TITLE

def add_h1(doc, num, text):
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(f'第{num}章　{text}')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = C_H1
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def add_h2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = C_H2
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

def add_h3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = C_H3
    p.paragraph_format.space_before = Pt(4)

def add_body(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))

def add_simple_table(doc, headers, rows, header_bg=(0x2E, 0x75, 0xB6)):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()

def add_risk_box(doc, level, title, content):
    color_map = {'HIGH': C_RED, 'MED': C_ORANGE, 'LOW': C_GREEN}
    label_map = {'HIGH': '[高風險]', 'MED': '[中風險]', 'LOW': '[改善建議]'}
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    hcell = tbl.rows[0].cells[0]
    hcell.text = f'{label_map[level]}  {title}'
    for run in hcell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = color_map[level]
        run.font.size = Pt(10.5)
    bcell = tbl.rows[1].cells[0]
    bcell.text = content
    for run in bcell.paragraphs[0].runs:
        run.font.size = Pt(10)
    doc.add_paragraph()

def add_code_block(doc, file_ref, code_lines):
    p = doc.add_paragraph()
    run = p.add_run(f'[程式碼佐證]  {file_ref}')
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = C_H1
    p.paragraph_format.space_after = Pt(0)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, (0xF2, 0xF2, 0xF2))
    cell.text = ''
    for i, line in enumerate(code_lines):
        p2 = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run2 = p2.add_run(line)
        run2.font.name = 'Courier New'
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
    doc.add_paragraph()

def add_toc(doc):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run('目　錄')
    run_t.font.size = Pt(16)
    run_t.font.bold = True
    run_t.font.color.rgb = C_TITLE
    p_title.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    r1 = p.add_run()
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    fc1.set(qn('w:dirty'), 'true')
    r1._r.append(fc1)
    r2 = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._r.append(instr)
    r3 = p.add_run()
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r3._r.append(fc2)
    r4 = p.add_run()
    r4.text = '（開啟後按 Ctrl+A → F9 更新目錄）'
    r4.font.size = Pt(9)
    r4.font.italic = True
    r4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r5 = p.add_run()
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r5._r.append(fc3)
    doc.add_page_break()

# ════════════════════════════════
# 封面
# ════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_title(doc, 'PPOST 出口進倉系統')
add_title(doc, '件數比對邏輯修正分析')
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('申報件數 vs. 報單明細件數 — 根因分析與修正說明')
run.font.size = Pt(13)
run.font.color.rgb = C_H2

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ('文件版本', 'v1.0'),
    ('建立日期', '2026-04-20'),
    ('涵蓋模組', 'TradeVan.POST.MnuGCIO / frmOpExpGCIO.cs'),
    ('分析人員', 'Pixiu AI Analysis'),
]
add_simple_table(doc, ['項目', '說明'], meta, header_bg=(0x1F, 0x49, 0x7D))
doc.add_page_break()

# ════════════════════════════════
# 目錄
# ════════════════════════════════
add_toc(doc)

# ════════════════════════════════
# 第一章：發生現象
# ════════════════════════════════
add_h1(doc, '一', '發生現象')
add_body(doc, '使用者回報截圖顯示：申報件數：3448 與 郵包貨物清單件數：3448 相同，系統卻報「件數比對不符」。')
add_body(doc, '兩個數值在畫面上顯示完全一致，但系統仍觸發硬擋錯誤，導致無法完成進倉確認作業。')

add_risk_box(doc, 'HIGH', '系統誤判導致無法進倉',
    '申報件數與清單件數畫面顯示相同，系統卻報「比對不符」並強制中斷，\n'
    '使用者無法繼續進倉操作，業務流程完全卡死。')

# ════════════════════════════════
# 第二章：系統中存在兩個比對點
# ════════════════════════════════
add_h1(doc, '二', '系統中存在兩個比對點')
add_body(doc, '系統在不同時間點各有一次件數比對，行為與影響完全不同。')

add_simple_table(doc, ['比對點', '位置', '觸發時機', '行為'],
[
    ('比對點 A', 'L455 Query()', '使用者輸入 HWB 查詢後', '詢問式（YesNo），可選擇繼續'),
    ('比對點 B', 'L1456 SaveEXE()', '按下「進倉確認」按鈕時', '強制中斷（return），無法繞過'),
])

add_h2(doc, '比對點 A — Query() 載入時（警示性詢問）')
add_body(doc, '觸發時機：使用者輸入 HWB 查詢後，系統自動執行。行為為詢問式（YesNo），使用者可選擇繼續。')
add_code_block(doc, 'frmOpExpGCIO.cs : L455-457', [
    'else if (Int32.Parse(MasterDataRow["PIECE"].ToString()) != _hwbDt.Rows.Count)',
    '{',
    '    DialogResult result = MessageBox.Show(String.Format(',
    '        "申報件數：{0}，郵包貨物清單件數：{1}，是否要人工確認進倉",',
    '        MasterDataRow["PIECE"].ToString(), _hwbDt.Rows.Count), "",',
    '        MessageBoxButtons.YesNo);',
    '}',
])

add_h2(doc, '比對點 B — SaveEXE() 確認進倉時（硬擋）')
add_body(doc, '觸發時機：使用者按下「進倉確認」按鈕時。行為為強制中斷（return），完全無法繞過。')
add_code_block(doc, 'frmOpExpGCIO.cs : L1456-1459', [
    '_DeclNoHwbDt = da.ReadByDeclNo(MasterDataRow["DECLNO"].ToString(), _whArea);',
    'if (Int32.Parse(MasterDataRow["PIECE"].ToString()) != _DeclNoHwbDt.Rows.Count)',
    '{',
    '    MessageBox.Show(String.Format(',
    '        "申報件數：{0}，此報單明細件數：{1}，件數比對不符，請人工調整確認",',
    '        MasterDataRow["PIECE"].ToString(), _DeclNoHwbDt.Rows.Count));',
    '    return;  // 直接中斷，無法進倉',
    '}',
])

# ════════════════════════════════
# 第三章：各欄位完整來源追蹤
# ════════════════════════════════
add_h1(doc, '三', '各欄位完整來源追蹤')

add_h2(doc, 'MasterDataRow（主資料列）')
add_simple_table(doc, ['項目', '說明'],
[
    ('來源方法', 'Query() L291：da.ReadByHWB(editHWB.Text.Trim(), _whArea)'),
    ('資料表', 'WORKEXPHWB'),
    ('SQL 條件', 'WHERE HWB = @HWB AND WHAREA = @WHAREA'),
    ('意義', '使用者輸入的那筆分提單（HWB）記錄'),
])

add_h2(doc, '{0} 申報件數 — MasterDataRow["PIECE"]')
add_simple_table(doc, ['項目', '說明'],
[
    ('資料表.欄位', 'WORKEXPHWB.PIECE'),
    ('意義', '此 HWB 所屬報單的申報總件數（寫在 HWB 記錄上）'),
    ('注意', '同一報單下所有 HWB 共用此值'),
])

add_h2(doc, '{1}（舊版）— _hwbDt.Rows.Count（顯示錯誤）')
add_simple_table(doc, ['項目', '說明'],
[
    ('來源方法', 'combDgvHWBList() L1693：da.ReadByMWB(mwb, _whArea)'),
    ('資料表', 'WORKEXPHWB'),
    ('SQL 條件', 'WHERE MWB = @MWB AND WHAREA = @WHAREA'),
    ('意義', '主提單（MWB）下的全部 HWB 筆數（跨報單）'),
    ('問題根因', '一個 MWB 可對應多張報單，此筆數範圍過大，與比對邏輯不一致'),
])

add_risk_box(doc, 'HIGH', '比對點 A 訊息框用錯資料來源',
    '比對點 A（L455）顯示的「郵包貨物清單件數」來自 _hwbDt（MWB 範圍），\n'
    '而實際硬擋比對（L1456）使用 _DeclNoHwbDt（DECLNO 範圍）。\n'
    '兩者範圍不同，導致畫面顯示數值與實際比對數值不一致，造成使用者誤解。')

add_h2(doc, '{1}（修正）— _DeclNoHwbDt.Rows.Count')
add_simple_table(doc, ['項目', '說明'],
[
    ('來源方法', 'SaveEXE() L1455：da.ReadByDeclNo(MasterDataRow["DECLNO"].ToString(), _whArea)'),
    ('資料表', 'WORKEXPHWB'),
    ('SQL 條件', 'WHERE DECLNO = @DECLNO AND WHAREA = @WHAREA'),
    ('意義', '同一報單（DECLNO）下的全部 HWB 筆數（同報單範圍）'),
    ('正確原因', '與比對邏輯使用相同的 DECLNO 範圍，數值一致'),
])

add_risk_box(doc, 'LOW', '建議：統一比對點 A 的訊息顯示',
    '將比對點 A（L455）的訊息框中「郵包貨物清單件數」改為使用 _DeclNoHwbDt.Rows.Count，\n'
    '與比對點 B 邏輯保持一致，避免使用者看到不一致的數值說明。')

# ════════════════════════════════
# 第四章：根因分析
# ════════════════════════════════
add_h1(doc, '四', '根因分析（Root Cause）')

add_h2(doc, '問題核心')
add_body(doc, '比對點 A（警告訊息）與比對點 B（硬擋邏輯）使用了不同的資料來源計算「件數」，導致：')
add_bullet(doc, '畫面警告顯示：申報件數 3448 vs 清單件數 3448 → 相同（但實際上是 MWB 範圍件數）')
add_bullet(doc, '進倉硬擋比對：申報件數 3448 vs DECLNO 範圍件數 N → 不同 → 觸發錯誤')

add_h2(doc, '關鍵程式碼')
add_code_block(doc, 'frmOpExpGCIO.cs : L1456-1458', [
    'if (Int32.Parse(MasterDataRow["PIECE"].ToString()) != _DeclNoHwbDt.Rows.Count)',
    '{',
    '    MessageBox.Show(String.Format("申報件數：{0}，此報單明細件數：{1}，件數比對不符...",',
    '        MasterDataRow["PIECE"].ToString(),',
    '        _DeclNoHwbDt.Rows.Count));',
    '}',
])

add_h2(doc, '資料來源驗證')
add_simple_table(doc, ['欄位', '來源', 'SQL 範圍', '說明'],
[
    ('申報件數', 'MasterDataRow["PIECE"]', '單筆 HWB 記錄', '報單定義的總件數，同報單共用'),
    ('此報單明細件數', '_DeclNoHwbDt.Rows.Count', 'WHERE DECLNO = @DECLNO', '資料庫中實際掛在該報單的 HWB 筆數'),
])

add_risk_box(doc, 'MED', '兩者不一致時的業務意義',
    '當 PIECE（申報） ≠ DECLNO 下的 HWB 筆數時，代表資料庫明細與申報主檔有落差，\n'
    '可能原因：部分 HWB 尚未建檔、多報單共用 MWB 造成分配問題、或資料匯入錯誤。\n'
    '需要人工確認後才能繼續進倉。')

doc.add_page_break()

# 儲存
out_path = os.path.join(DOCS_DIR, '2026-04-20-115554-PPOST-件數比對邏輯修正.docx')
doc.save(out_path)
print(f'DOCX saved: {out_path}')
