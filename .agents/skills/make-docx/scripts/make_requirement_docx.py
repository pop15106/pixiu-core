# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, size=9.5, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_rows=1):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=9.5)
            if row_idx < header_rows:
                set_cell_shading(cell, "1F4E79")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, "EAF2F8")


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.text = ""
    run = p.add_run(text)
    set_run_font(
        run,
        size=15 if level == 1 else 12.5,
        bold=True,
        color=(31, 78, 121) if level == 1 else (55, 55, 55),
    )
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(str(item))
        set_run_font(run, size=10.5)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(str(item))
        set_run_font(run, size=10.5)
        p.paragraph_format.space_after = Pt(3)


def add_kv_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    for idx, row in enumerate(rows):
        key, value = row[0], row[1]
        set_cell_text(table.cell(idx, 0), key, bold=True)
        set_cell_shading(table.cell(idx, 0), "D9EAF7")
        set_cell_text(table.cell(idx, 1), value)
    style_table(table, header_rows=0)


def add_matrix_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table, header_rows=1)


def ftp_monthly_message_report_spec():
    return {
        "title": "FTP 訊息單量彙整自動化需求規格書",
        "subtitle": "需求規格與可實作方案評估",
        "meta": [
            ["文件版本", "v1.0"],
            ["文件日期", "2026-04-27"],
            ["需求主題", "每月 FTP 訊息下載、單量確認與 Excel 月報自動化"],
            ["使用對象", "承辦人員、主管、開發團隊"],
            ["文件目的", "將人工月報流程轉為可評估、可開發、可驗收的需求規格"],
        ],
        "summary": "本文件整理每月人工至 FTP 下載訊息、逐檔確認單量並同步 Excel 的作業流程，轉換為可實作的自動化需求規格，並提供多種落地方案供主管與開發團隊評估。",
        "sections": [
            {
                "heading": "一、需求背景與目標",
                "paragraphs": ["目前每月需人工登入 FTP，下載當月所有訊息檔，逐一開啟檔案確認單量，再整理至 Excel 給主管。此流程耗時、容易漏檔或重複統計，也不易追溯每個數字的來源。"],
                "tables": [{
                    "headers": ["項目", "說明"],
                    "rows": [
                        ["現況痛點", "人工下載、人工開檔、人工統計、人工填 Excel，耗時且容易出錯。"],
                        ["改善目標", "自動抓取指定月份 FTP 訊息檔，解析並統計單量，輸出固定格式 Excel。"],
                        ["成功標準", "統計結果可追溯來源檔案，可重跑、可檢核、可交付主管。"],
                        ["優先原則", "先做最小可行版本，再逐步擴充排程、通知與系統整合。"],
                    ],
                }],
            },
            {
                "heading": "二、功能需求",
                "tables": [{
                    "headers": ["編號", "功能", "需求說明", "驗收條件"],
                    "rows": [
                        ["FR-01", "指定月份", "使用者可輸入年月，預設為當月。", "可執行當月與歷史月份資料整理。"],
                        ["FR-02", "FTP 檔案取得", "系統連線 FTP 並列出指定路徑檔案。", "FTP 檔案數與下載清單一致。"],
                        ["FR-03", "月份篩選", "依檔名、資料夾或內容日期篩選當月訊息。", "非本月檔案不納入統計。"],
                        ["FR-04", "訊息解析", "逐檔讀取訊息內容，取得單量與分類欄位。", "解析結果可回溯原始檔名。"],
                        ["FR-05", "單量統計", "依日期、訊息類型或主管需求彙總件數。", "總數與明細加總一致。"],
                        ["FR-06", "Excel 輸出", "將統計結果寫入固定格式 Excel。", "主管可直接開啟查看。"],
                        ["FR-07", "異常處理", "記錄下載失敗、格式錯誤、缺欄位等異常。", "異常檔案會列在獨立清單。"],
                        ["FR-08", "重跑保護", "同月份重跑不應重複累加。", "重新產生或覆蓋該月份報表。"],
                    ],
                }],
            },
            {
                "heading": "三、可實作方案比較",
                "tables": [{
                    "headers": ["方案", "技術組合", "優點", "代價 / 風險", "適用情境", "建議度"],
                    "rows": [
                        ["A", "PowerShell + Excel", "Windows 原生、部署簡單、適合快速 MVP。", "解析複雜格式較吃力；Excel 被開啟時可能寫入失敗。", "訊息格式簡單，只供少數人使用。", "中"],
                        ["B", "Python Batch + openpyxl", "解析彈性高、Excel 寫入穩定、容易加 Log 與測試。", "需處理 Python 環境或打包成 exe。", "最適合第一版正式自動化。", "高"],
                        ["C", "Java Batch", "可與既有 Java 系統治理一致，適合正式排程。", "開發成本較高，可能牽動既有部署流程。", "未來要多人使用或納入正式系統。", "中高"],
                        ["D", "Power Automate / RPA", "貼近人工操作流程，上手快。", "授權與穩定性需確認；畫面或格式變動易壞。", "公司已有 RPA 工具且不便寫程式。", "中"],
                    ],
                }],
            },
            {
                "heading": "四、驗收測試建議",
                "tables": [{
                    "headers": ["測試案例", "測試內容", "預期結果"],
                    "rows": [
                        ["正常月份", "指定一個有完整檔案的月份執行。", "成功產出 Excel，總數正確。"],
                        ["無檔月份", "指定沒有訊息檔的月份。", "產出空報表或明確提示無資料。"],
                        ["格式錯誤檔", "放入一筆無法解析的訊息。", "不中斷整批作業，異常清單列出該檔。"],
                        ["重複執行", "同月份連續執行兩次。", "不重複累加，結果一致。"],
                    ],
                }],
            },
        ],
    }


def load_spec(args):
    if args.sample == "ftp-monthly-message-report":
        return ftp_monthly_message_report_spec()
    if args.spec_json:
        with Path(args.spec_json).open("r", encoding="utf-8") as fh:
            return json.load(fh)
    raise SystemExit("請提供 --spec-json 或 --sample ftp-monthly-message-report")


def render_docx(spec, output):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "List Bullet", "List Number"]:
        if style_name in styles:
            styles[style_name].font.name = "Microsoft JhengHei"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(spec.get("title", "需求規格書"))
    set_run_font(run, size=22, bold=True, color=(31, 78, 121))

    subtitle = spec.get("subtitle")
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        set_run_font(run, size=14, bold=True, color=(91, 91, 91))
        p.paragraph_format.space_after = Pt(24)

    add_kv_table(doc, spec.get("meta", []))

    summary = spec.get("summary")
    if summary:
        add_paragraph(doc, "文件摘要", bold=True)
        add_paragraph(doc, summary)

    if spec.get("sections"):
        doc.add_page_break()

    for section_spec in spec.get("sections", []):
        add_heading(doc, section_spec["heading"])
        for text in section_spec.get("paragraphs", []):
            add_paragraph(doc, text)
        for item in section_spec.get("bullets", []):
            add_bullets(doc, [item])
        for item in section_spec.get("numbered", []):
            add_numbered(doc, [item])
        for table in section_spec.get("tables", []):
            add_matrix_table(doc, table["headers"], table["rows"])

    for doc_section in doc.sections:
        footer = doc_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"{spec.get('title', '需求規格書')} ｜ 內部評估文件")
        set_run_font(run, size=9, color=(120, 120, 120))

    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    if output.exists():
        stem = output.stem
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output = output.with_name(f"{stem}_{stamp}{output.suffix}")
    doc.save(output)
    return output


def main():
    parser = argparse.ArgumentParser(description="產生 Pixiu 風格需求規格 DOCX。")
    parser.add_argument("--spec-json", help="UTF-8 JSON 規格檔路徑。")
    parser.add_argument("--sample", choices=["ftp-monthly-message-report"], help="使用內建範例。")
    parser.add_argument("--output", required=True, help="輸出 .docx 路徑。")
    args = parser.parse_args()

    spec = load_spec(args)
    output = render_docx(spec, args.output)
    print(f"DOCX saved: {output}")


if __name__ == "__main__":
    main()
